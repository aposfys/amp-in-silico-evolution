"""
Comprehensive Word document: MACREL+ESM3 vs MACREL+ESM3+PFES
All graphs embedded, full analytical explanations.
"""
import os, warnings
import pandas as pd
import numpy as np
warnings.filterwarnings('ignore')

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── paths ─────────────────────────────────────────────────────────────────────
LOG_A = "results_macrel_300/progress.log"
LOG_B = "results_macrel_pfes_300/progress.log"
DIR_A = "results_macrel_300/analysis"
DIR_B = "results_macrel_pfes_300/analysis"
CDIR  = "comparison_plots_300"
OUT   = "AMP_evolution_analysis.docx"

LABEL_A = "MACREL + ESM3"
LABEL_B = "MACREL + ESM3 + PFES"

# ── load data for metrics ─────────────────────────────────────────────────────
def load(path):
    df = pd.read_csv(path, sep='\t', comment='#', low_memory=False)
    amp = 'amp_prob' if 'amp_prob' in df.columns else 's_amp'
    df['amp'] = pd.to_numeric(df[amp], errors='coerce')
    for c in ['score','hemo_prob','mean_plddt','ptm','seq_len']:
        df[c] = pd.to_numeric(df[c], errors='coerce')
    df['gen'] = df['gndx'].str.extract(r'(\d+)').astype(float)
    return df.dropna(subset=['score','gen'])

A, B = load(LOG_A), load(LOG_B)

# compute key metrics
def run_metrics(df, label):
    best_row = df.loc[df['score'].idxmax()]
    fin = df[df.gen == df.gen.max()]
    return {
        'label': label,
        'best_score': df['score'].max(),
        'best_gen': int(best_row['gen']),
        'best_seq': best_row.get('sequence','N/A'),
        'best_seq_len': int(best_row['seq_len']) if not pd.isna(best_row['seq_len']) else '?',
        'best_plddt': best_row['mean_plddt'],
        'best_ptm': best_row['ptm'],
        'best_amp': best_row['amp'],
        'best_hemo': best_row['hemo_prob'],
        'final_mean_score': fin['score'].mean(),
        'final_mean_len': fin['seq_len'].mean(),
        'final_mean_amp': fin['amp'].mean(),
        'final_mean_hemo': fin['hemo_prob'].mean(),
        'final_mean_plddt': fin['mean_plddt'].mean(),
        'total_seqs': len(df),
    }

mA = run_metrics(A, LABEL_A)
mB = run_metrics(B, LABEL_B)

# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT SETUP
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

# Page margins (2.5 cm each side)
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── helpers ───────────────────────────────────────────────────────────────────
def set_font(run, size=11, bold=False, italic=False,
             color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x37, 0x5A)
        run.font.size = Pt(18)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
        run.font.size = Pt(14)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
        run.font.size = Pt(12)
    return p

def body(text, bold=False, italic=False, space_after=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    if space_after:
        p.paragraph_format.space_after = Pt(6)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.3 + level*0.25)
    p.paragraph_format.space_after = Pt(3)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x60)
    return p

def add_image(path, caption, width=5.8):
    if not os.path.exists(path):
        body(f"[Image not found: {path}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.size = Pt(9.5)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    cap.paragraph_format.space_after = Pt(10)

def separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1A375A')
        tcPr.append(shd)
    # data rows
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri+1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)
            for run in cells[ci].paragraphs[0].runs:
                run.font.size = Pt(10)
            if ri % 2 == 0:
                tc = cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F4F6F9')
                tcPr.append(shd)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Investigating Novel Antimicrobial Peptides\nUsing Machine Learning, Structure Prediction\nand In Silico Evolution")
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x37, 0x5A)
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Comparative Analysis: MACREL+ESM3  vs  MACREL+ESM3+PFES")
r.font.size = Pt(14)
r.font.italic = True
r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Apostolos Fysekidis\nMSc Bioinformatics – Computational Biology\nNational & Kapodistrian University of Athens\n2025")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Project Overview")

body("This document reports the results of two independent in silico directed evolution experiments for antimicrobial peptide (AMP) discovery. Both experiments use the same evolutionary framework (PFES) and the same structure predictor (ESM3), but differ in the fitness function applied to evaluate and select sequences. The goal is to understand how the choice of fitness function shapes the evolutionary trajectory, the properties of the final sequences, and whether the resulting candidates are realistic drug leads.")

body("The two variants are:")
bullet("MACREL + ESM3: fitness = pLDDT × pTM × AMP_prob × (1 − hemo_proxy)  [4 terms]")
bullet("MACREL + ESM3 + PFES: fitness = pLDDT × pTM × length_penalty × helix_penalty × β_penalty × AMP_prob × (1 − hemo_proxy) × contact_density  [8 terms]")

body("Both runs used 300 generations, a population of 100 sequences, and one ESM3 structure prediction per sequence per generation — totalling 30,000 predictions per run.")

separator()

# ══════════════════════════════════════════════════════════════════════════════
# 2. THE EVOLUTIONARY LOOP
# ══════════════════════════════════════════════════════════════════════════════
h1("2. The Evolutionary Framework (PFES)")

add_image("loop_diagram.png",
    "Figure 1. The in silico evolution loop. Each generation: mutate all sequences, predict structure with ESM3, score with the fitness function, select 100 survivors via Boltzmann sampling.",
    width=5.8)

body("The pipeline is based on PFES (Protein Fitness Evolutionary Search, Sahakyanhk et al., PNAS 2022), adapted for AMP discovery. Each generation consists of four steps:")

h2("2.1 MUTATE")
body("Every sequence in the population (100 sequences) is mutated exactly once. Five operators act with empirically calibrated probabilities:")
add_table(
    ["Operator", "Symbol", "Approx. probability", "Effect"],
    [
        ["Point substitution", "~", "~42%", "Replace one residue with any amino acid"],
        ["Insertion", "+", "~4%", "Insert one residue at a random position"],
        ["Deletion", "−", "~4%", "Remove one residue"],
        ["Scramble", "*", "~1.7%", "Shuffle a random segment"],
        ["Duplication", "d", "~0.2%", "Duplicate a segment (rare but consequential)"],
    ]
)
body("The duplication operator is rare, but critical: it can double the length of a sequence in one step. Without a length penalty in the fitness function, this operator is strongly rewarded because longer sequences accumulate more pLDDT contacts and higher apparent pTM scores.", italic=True)

h2("2.2 PREDICT STRUCTURE — ESM3")
body("ESM3 (EvolutionaryScale, 2024) folds each mutated sequence using a protein language model. It returns:")
bullet("pLDDT (per-residue local distance difference test): values in [0, 1]. >0.9 = confident, well-structured residue; <0.5 = likely disordered.")
bullet("pTM (predicted Template Modelling score): values in [0, 1]. >0.5 = plausible global fold topology. For peptides <30 aa, values of 0.5–0.6 are normal and do not indicate poor folding.")
bullet("Backbone 3D coordinates: used to compute secondary structure and Cβ–Cβ contact density.")
body("The original PFES paper (Sahakyanhk 2022) used ESMFold v1. This project replaces it with ESM3, which is compatible with Apple MPS and provides equivalent predictions for short peptides.")

h2("2.3 SCORE")
body("Each sequence receives a multiplicative fitness score. Because all terms are multiplied together, a sequence must score well on every criterion simultaneously — a single near-zero term collapses the entire score. This creates selection pressure toward sequences that are globally competent, rather than extremely good at one property.")

h2("2.4 SELECT — Boltzmann Sampling")
body("100 survivors are drawn from the combined pool (100 old + 100 mutated = 200 sequences) using Boltzmann-weighted sampling:")
code_block("P(sequence i) = exp(β × score_i) / Σ exp(β × score_j)        β = 20")
body("With β = 20, a sequence scoring 0.65 is approximately 55× more likely to survive than one scoring 0.50. This creates strong selection pressure toward high-scoring sequences while preserving diversity — unlike deterministic top-N selection, some lower-scoring sequences always survive, which helps escape local optima.")

body("Key numbers per run:")
bullet("300 generations × 100 mutations = 30,000 ESM3 predictions")
bullet("The --norepeat flag ensures no sequence is evaluated twice across the entire run")
bullet("Starting sequences: random 24 aa amino acid strings")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. TOOLS & SCORING COMPONENTS
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Scoring Components — What Each Term Measures")

h2("3.1 MACREL — AMP Classifier")
body("MACREL (Santos-Júnior et al., PeerJ 2020) is a machine learning classifier for antimicrobial peptide prediction. It uses two ONNX random forest models trained on experimentally validated AMPs:")
bullet("AMP probability: probability that a sequence is antimicrobial [0, 1]. Threshold for prediction: 0.5.")
bullet("Hemolytic probability: probability that a sequence lyses red blood cells. However, MACREL's hemolytic model was trained on only 220 peptides (Chaudhary 2016). Evolved sequences — which are highly cationic (charge > +6) and K/R-dominant — fall outside this training distribution. The model outputs exactly 0.000 for all evolved sequences, making it useless as a selection signal.")
body("MACREL is called on batches of sequences as a subprocess, parsing its TSV output. It processes full protein sequences without length restriction.")

body("MACREL saturation: AMP probability exceeds 0.88 in both runs by generation ~80 and then plateaus. After saturation, the AMP term contributes a near-constant multiplier (~0.90) and stops driving selection. All further evolution is driven by structural terms and the hemolytic proxy.", bold=True)

h2("3.2 ESM3 — Structure Predictor")
body("ESM3 is a protein language model that performs structure prediction from sequence. It produces:")
bullet("pLDDT: per-residue confidence. Mean pLDDT is used as a fitness term.")
bullet("pTM: predicted TM-score against PDB templates. Used as a global fold quality term.")
bullet("3D backbone coordinates: used downstream to compute contacts and secondary structure.")

h2("3.3 Hemolytic Proxy (Biophysical)")
body("Since MACREL's hemolytic model fails for evolved sequences, a biophysical proxy was implemented:")
code_block("hemo_proxy = sigmoid( hydrophobicity_ratio × 10 − min(charge_pH7.4, 8.0) × 0.5 − 2.0 )")
body("This formula captures the two main drivers of hemolysis:")
bullet("Hydrophobicity: hydrophobic peptides insert into red blood cell membranes. Higher hydrophobicity → higher hemo risk.")
bullet("Positive charge: cationic peptides are electrostatically rejected by the zwitterionic outer leaflet of RBCs, even if hydrophobic. Higher charge → lower hemo risk.")
body("The charge is capped at +8 to prevent the optimizer from driving charge to extreme values simply to obtain a free hemo pass. Calibration against reference peptides:")
add_table(
    ["Peptide", "Known hemolysis", "Proxy score", "Assessment"],
    [
        ["Mastoparan (wasp toxin)", "Strongly hemolytic", "0.97", "Correct ✓"],
        ["Melittin (bee venom)", "Moderately hemolytic", "0.62", "Correct ✓"],
        ["Magainin-2 (frog AMP)", "Non-hemolytic", "0.70", "Fails ✗"],
        ["Best PFES sequence", "Expected non-hemolytic", "0.15", "Plausible ✓"],
    ]
)
body("The proxy correctly identifies highly hemolytic peptides but is imperfect for borderline cases like magainin-2. Experimental validation is required.", italic=True)

h2("3.4 PFES Structural Penalties (MACREL+ESM3+PFES only)")
body("Four additional penalty terms were taken from the PFES framework:")

body("Length penalty — penalises sequences longer than 30 aa:")
code_block("length_penalty = 1 − sigmoid(seq_len, threshold=30, slope=0.2)")
body("A 26 aa sequence scores ~0.69; a 50 aa sequence scores ~0.08; an 80 aa sequence scores ~0.001. This prevents the duplication operator from exploiting score inflation.")

body("Helix penalty — rewards sequences with alpha-helical content:")
code_block("helix_penalty = 1 − exp(−helix_fraction × 3)")
body("Peptides with high helix content (>40%) approach penalty ≈ 1.0 (no penalty). Fully disordered sequences are penalised.")

body("Beta-sheet penalty — penalises sequences with beta-sheet content:")
code_block("beta_penalty = exp(−beta_fraction × 5)")
body("AMPs are alpha-helical; beta sheets correlate with aggregation. This term acts as a penalty.")

body("Contact density — rewards compact structures:")
code_block("contact_density = (Cβ–Cβ contacts within 6Å + seq_len) / seq_len")
body("This is the only term that can exceed 1.0. For a 26 aa helical peptide, contact density ≈ 2.0–2.5. It is essential for short peptides to remain competitive: without it, short sequences have fewer total contacts and lower absolute pLDDT sums, making them disadvantaged compared to longer sequences.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. RUN A: MACREL + ESM3
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Run A — MACREL + ESM3")

body("Score formula:")
code_block("score = pLDDT × pTM × AMP_prob × (1 − hemo_proxy)")

body("This is the baseline run. The four-term fitness function rewards confident structure (pLDDT, pTM) and antimicrobial activity (AMP_prob) while weakly penalising hemolysis (hemo_proxy). There is no constraint on sequence length, secondary structure, or compactness.")

h2("4.1 Key Metrics")
add_table(
    ["Metric", "Value"],
    [
        ["Best fitness score", f"{mA['best_score']:.4f}"],
        ["Best score at generation", f"{mA['best_gen']}"],
        ["Best sequence", f"{mA['best_seq'][:50]}{'...' if len(str(mA['best_seq']))>50 else ''}"],
        ["Best sequence length", f"{mA['best_seq_len']} aa"],
        ["Best pLDDT", f"{mA['best_plddt']:.4f}"],
        ["Best pTM", f"{mA['best_ptm']:.4f}"],
        ["Best AMP probability", f"{mA['best_amp']:.4f}"],
        ["Best hemolytic proxy", f"{mA['best_hemo']:.4f}"],
        ["Final generation mean score", f"{mA['final_mean_score']:.4f}"],
        ["Final generation mean length", f"{mA['final_mean_len']:.1f} aa"],
        ["Final generation mean AMP prob", f"{mA['final_mean_amp']:.4f}"],
        ["Total sequences evaluated", f"{mA['total_seqs']:,}"],
    ]
)

h2("4.2 Fitness Evolution")
add_image(f"{DIR_A}/Evolution.png",
    "Figure 2. Fitness score over 300 generations — MACREL+ESM3. Solid line: best score per generation. Dashed: population mean. Shaded band: IQR.",
    width=5.8)

body("Interpretation:")
bullet("Generations 0–80: rapid fitness gain driven by accumulation of the MACREL AMP signal. Random sequences quickly converge to cationic amphipathic motifs that score well on AMP probability.")
bullet("Generations 80–300: plateau. AMP probability has saturated (>0.90 for the entire population). Only pLDDT, pTM, and the hemolytic proxy continue to drive selection.")
bullet("Best score: 0.626 at generation ~245. With pLDDT ≈ 0.99, pTM ≈ 0.84, AMP ≈ 0.90, hemo ≈ 0.17: 0.99 × 0.84 × 0.90 × 0.83 ≈ 0.62, consistent with observed values.")

h2("4.3 Score Components")
add_image(f"{DIR_A}/Score_components.png",
    "Figure 3. Individual score component trajectories — MACREL+ESM3.",
    width=5.8)

body("Key observations:")
bullet("pLDDT rises rapidly to >0.98 and stays there. Helix-forming sequences have near-perfect per-residue confidence in ESM3.")
bullet("pTM climbs from ~0.4 to ~0.85. Higher for longer sequences — pTM rewards structural complexity that matches PDB templates.")
bullet("AMP probability: rapid rise to ~0.88 within 80 generations, then flat. Saturation is the dominant feature.")
bullet("Hemolytic proxy: stays low (~0.10–0.18) throughout. The charge term in the proxy suppresses hemolysis risk even as hydrophobicity grows.")

h2("4.4 Fitness Landscape")
add_image(f"{DIR_A}/Fitness_landscape.png",
    "Figure 4. Fitness landscape — MACREL+ESM3. Score vs. sequence diversity.",
    width=5.8)

h2("4.5 Amino Acid Composition")
add_image(f"{DIR_A}/AA_composition.png",
    "Figure 5. Amino acid frequency in final 10 generations — MACREL+ESM3.",
    width=5.8)

body("Strongly enriched residues:")
bullet("K (Lysine) ~25%, R (Arginine) ~22%: cationic residues that drive electrostatic binding to the negatively charged bacterial outer membrane.")
bullet("M (Methionine) ~12%, L (Leucine) ~10%: this is a duplication artefact. The optimizer duplicated a short M/L-rich motif multiple times, growing sequences to 80 aa. Both are hydrophobic but their enrichment beyond ~8% marks the length artefact.")

body("Depleted residues:")
bullet("D, E: anionic — eliminated in both runs because they penalise membrane binding and raise the hemolytic proxy.")
bullet("G, P: helix breakers — their absence reinforces alpha-helical structure.")

h2("4.6 Score Distribution (Final Generation)")
add_image(f"{DIR_A}/Score_distribution.png",
    "Figure 6. Distribution of fitness scores in generation 299 — MACREL+ESM3.",
    width=5.5)

body(f"Final generation (gen 299): mean = {mA['final_mean_score']:.3f}. The distribution is wide (range ~0.35–0.63), reflecting heterogeneous sequence lengths within the population — sequences that duplicated more aggressively score higher due to length inflation.")

h2("4.7 Secondary Structure")
add_image(f"{DIR_A}/Secondary_structures.png",
    "Figure 7. Secondary structure content over generations — MACREL+ESM3.",
    width=5.8)

body("Alpha-helix content rises to >70% by generation 50 and stabilises. The fitness function implicitly rewards helical structure: pLDDT is highest for well-defined secondary structure elements, and cationic amphipathic helices are the dominant AMP structural motif.")

h2("4.8 Run Summary")
add_image(f"{DIR_A}/Summary.png",
    "Figure 8. Summary panel — MACREL+ESM3.",
    width=5.8)

body("The run effectively discovers AMP-like sequences from random starts within 80 generations. However, the optimizer exploits the duplication operator to grow sequences to 80 aa. The result is sequences that score well on the fitness function but are pharmacologically unrealistic: synthesis cost scales steeply with length (>€1,000/mg at 80 aa vs ~€150/mg at 26 aa), and longer peptides face increased proteolysis risk and potential immunogenicity.", bold=False)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. RUN B: MACREL + ESM3 + PFES
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Run B — MACREL + ESM3 + PFES")

body("Score formula:")
code_block("score = pLDDT × pTM × length_penalty × helix_penalty × β_penalty\n        × AMP_prob × (1 − hemo_proxy) × contact_density")

body("This run adds four structural constraints from the PFES framework to the four MACREL+ESM3 terms. The length penalty prevents sequence growth beyond 30 aa; the helix/β penalties reward alpha-helical and penalise beta-sheet content; contact density rewards compact, well-packed structures. The optimizer must simultaneously satisfy all eight constraints to achieve a high score.")

h2("5.1 Key Metrics")
add_table(
    ["Metric", "Value"],
    [
        ["Best fitness score", f"{mB['best_score']:.4f}"],
        ["Best score at generation", f"{mB['best_gen']}"],
        ["Best sequence", f"{mB['best_seq'][:50]}{'...' if len(str(mB['best_seq']))>50 else ''}"],
        ["Best sequence length", f"{mB['best_seq_len']} aa"],
        ["Best pLDDT", f"{mB['best_plddt']:.4f}"],
        ["Best pTM", f"{mB['best_ptm']:.4f}"],
        ["Best AMP probability", f"{mB['best_amp']:.4f}"],
        ["Best hemolytic proxy", f"{mB['best_hemo']:.4f}"],
        ["Final generation mean score", f"{mB['final_mean_score']:.4f}"],
        ["Final generation mean length", f"{mB['final_mean_len']:.1f} aa"],
        ["Final generation mean AMP prob", f"{mB['final_mean_amp']:.4f}"],
        ["Total sequences evaluated", f"{mB['total_seqs']:,}"],
    ]
)

h2("5.2 Fitness Evolution")
add_image(f"{DIR_B}/Evolution.png",
    "Figure 9. Fitness score over 300 generations — MACREL+ESM3+PFES.",
    width=5.8)

body("Interpretation:")
bullet("Generations 0–80: same MACREL-driven rapid rise as Run A.")
bullet("Generations 80–150: slower convergence than Run A — the optimizer must satisfy 8 simultaneous constraints rather than 4.")
bullet("Peak score: 0.595. Lower than Run A (0.626) but the score is an honest reflection of all 8 constraints — there is no length inflation.")
bullet("Score breakdown for best sequence (26 aa): pLDDT 0.990 × pTM 0.580 × length_penalty 0.690 × helix_penalty ~1.0 × β_penalty 1.0 × AMP_prob 0.882 × (1−0.146) × contact_density ~2.0 ≈ 0.595")

h2("5.3 Score Components")
add_image(f"{DIR_B}/Score_components.png",
    "Figure 10. Individual score component trajectories — MACREL+ESM3+PFES.",
    width=5.8)

body("Key differences from Run A:")
bullet("pTM is lower (0.55–0.58 vs 0.82–0.85): short peptides (~26 aa) have less structure to align to PDB templates. This is normal for therapeutic-length peptides and does not indicate poor folding.")
bullet("Contact density term acts as a compensatory multiplier: for a compact 26 aa helix, contact density ≈ 2.0, partially offsetting the lower pTM.")
bullet("Length penalty stabilises around 0.69–0.72, consistent with ~26 aa sequences in the therapeutic window.")

h2("5.4 Fitness Landscape")
add_image(f"{DIR_B}/Fitness_landscape.png",
    "Figure 11. Fitness landscape — MACREL+ESM3+PFES.",
    width=5.8)

h2("5.5 Amino Acid Composition")
add_image(f"{DIR_B}/AA_composition.png",
    "Figure 12. Amino acid frequency in final 10 generations — MACREL+ESM3+PFES.",
    width=5.8)

body("Key differences from Run A:")
bullet("F (Phenylalanine) ~8% vs ~4% in Run A: aromatic hydrophobic residue enriched on the hydrophobic face of the amphipathic helix.")
bullet("No methionine explosion: length penalty prevents duplication from accumulating M-rich repeats. M stays at ~3%.")
bullet("K and R still dominant (~20–22% each) but at more moderate levels than Run A (~25–30%).")
bullet("Composition is more consistent with natural therapeutic AMPs: magainin-2 has multiple F residues; LL-37 (37 aa) is F-enriched.")

h2("5.6 Score Distribution (Final Generation)")
add_image(f"{DIR_B}/Score_distribution.png",
    "Figure 13. Distribution of fitness scores in generation 299 — MACREL+ESM3+PFES.",
    width=5.5)

body(f"Final generation (gen 299): mean = {mB['final_mean_score']:.3f}. The distribution is tighter than Run A. All sequences are constrained to similar lengths (~26 aa), so score variation reflects genuine fitness differences across the 8-dimensional constraint space rather than length heterogeneity.")

h2("5.7 Secondary Structure")
add_image(f"{DIR_B}/Secondary_structures.png",
    "Figure 14. Secondary structure content over generations — MACREL+ESM3+PFES.",
    width=5.8)

body("Alpha-helix content also dominates, consistent with Run A. The explicit helix_penalty term reinforces this — sequences with low helix content cannot score well on this term.")

h2("5.8 Run Summary")
add_image(f"{DIR_B}/Summary.png",
    "Figure 15. Summary panel — MACREL+ESM3+PFES.",
    width=5.8)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. COMPARISON
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Comparative Analysis")

h2("6.1 Fitness Trajectories")
add_image(f"{CDIR}/c_fitness.png",
    "Figure 16. Fitness over 300 generations — both runs overlaid. Blue: MACREL+ESM3. Green: MACREL+ESM3+PFES.",
    width=5.8)

body("Both runs follow the same trajectory shape: rapid gain in generations 0–80 (MACREL AMP signal), then plateau. MACREL+ESM3 achieves a higher absolute score (0.626 vs 0.595), but this is a length artefact — longer sequences accumulate more Cβ contacts and higher pTM, inflating all multiplicative terms. The lower score of MACREL+ESM3+PFES is more informative and biologically honest.")

h2("6.2 Sequence Length — The Defining Result")
add_image(f"{CDIR}/c_length.png",
    "Figure 17. Sequence length over 300 generations — both runs. Dashed reference: 30 aa therapeutic target.",
    width=5.8)

body("This is the single most important difference between the two runs:")
bullet("MACREL+ESM3: sequences grow continuously from 24 aa → 80 aa. The duplication operator (d) copies a well-scoring motif, and without a length penalty the optimizer is rewarded for doing so. Final population mean length: ~62 aa.")
bullet("MACREL+ESM3+PFES: sequences stabilise at 26 aa from generation 20 onward. The length penalty prevents any gain from growing beyond 30 aa.")

body("Clinical and pharmacological relevance:")
add_table(
    ["Peptide", "Length", "Status"],
    [
        ["Magainin-2", "23 aa", "Natural AMP, antifungal activity"],
        ["Pexiganan (MSI-78)", "22 aa", "Magainin analogue, Phase III trials"],
        ["LL-37", "37 aa", "Human cathelicidin, broad-spectrum"],
        ["Polymyxin B", "10 aa", "Clinical antibiotic, cyclic"],
        ["Best MACREL+ESM3 sequence", "80 aa", "Outside therapeutic window ✗"],
        ["Best MACREL+ESM3+PFES sequence", "26 aa", "Within therapeutic window ✓"],
    ]
)
body("Synthesis cost: ~€150/mg at 26 aa; >€1,000/mg at 80 aa. Longer peptides are also more susceptible to proteolysis and potentially immunogenic.", italic=True)

h2("6.3 AMP Probability — Saturation")
add_image(f"{CDIR}/c_amp.png",
    "Figure 18. AMP probability over 300 generations — both runs.",
    width=5.8)

body("AMP probability saturates in both runs by generation ~80, exceeding 0.88 in the best sequence and converging in the population mean. This is a fundamental limitation of using a binary ML classifier as a fitness signal: once the entire population satisfies the classifier's threshold, the signal provides no further discrimination. A quantitative MIC predictor (e.g. LLAMP, outputting log₁₀(MIC) in μM) would remain informative throughout all 300 generations and is the planned upgrade.")

h2("6.4 Hemolytic Proxy")
add_image(f"{CDIR}/c_hemo.png",
    "Figure 19. Hemolytic proxy over 300 generations — both runs.",
    width=5.8)

body("Both runs maintain hemo_proxy < 0.20 throughout. The proxy is driven by the balance between hydrophobicity (increases risk) and positive charge (decreases risk). R+K content rises to ~40–46% in both runs, pushing the charge term high enough to suppress the proxy. The charge cap at +8 prevents the optimizer from exploiting infinite charge to eliminate the hemo term entirely.")

h2("6.5 pLDDT")
add_image(f"{CDIR}/c_plddt.png",
    "Figure 20. Mean pLDDT over 300 generations — both runs.",
    width=5.8)

body("Both runs rapidly converge to pLDDT > 0.97. Alpha-helical peptides are structurally regular and ESM3 is confident about their residue positions. The slight edge for MACREL+ESM3 (reaching 0.99 vs 0.97) reflects the longer sequences — more residues embedded in a regular helix have higher per-residue confidence.")

h2("6.6 Amino Acid Composition Comparison")
add_image(f"{CDIR}/c_aa.png",
    "Figure 21. Amino acid frequency — final 10 generations of both runs side by side.",
    width=5.8)

add_table(
    ["Residue", "MACREL+ESM3", "MACREL+ESM3+PFES", "Interpretation"],
    [
        ["K (Lys)", "~25%", "~22%", "Cationic — enriched in both"],
        ["R (Arg)", "~22%", "~20%", "Cationic — enriched in both"],
        ["M (Met)", "~12%", "~3%", "Duplication artefact in Run A"],
        ["L (Leu)", "~10%", "~7%", "Hydrophobic, duplicated in Run A"],
        ["F (Phe)", "~4%", "~8%", "Aromatic hydrophobic — Run B enriched"],
        ["I (Ile)", "~3%", "~5%", "Branched hydrophobic — slightly enriched in Run B"],
        ["D, E", "~0%", "~0%", "Anionic — eliminated in both"],
        ["G, P", "~1%", "~1%", "Helix breakers — depleted in both"],
    ]
)

h2("6.7 Score Distribution Comparison")
add_image(f"{CDIR}/c_dist.png",
    "Figure 22. Fitness score distribution in final generation — both runs overlaid.",
    width=5.5)

body("MACREL+ESM3: mean 0.585, wide distribution (range 0.35–0.63). Spread reflects heterogeneous sequence lengths — different duplication histories within the population produce different score levels.")
body("MACREL+ESM3+PFES: mean 0.521, tighter distribution. All sequences are constrained to similar lengths, so variation reflects genuine multi-dimensional fitness differences.")
body("A tighter distribution at a lower mean is a sign of better-behaved optimization, not a worse result. The optimizer has converged to a consistent region of sequence space.", bold=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. SUMMARY TABLE
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Summary of Results")

add_table(
    ["Metric", "MACREL + ESM3", "MACREL + ESM3 + PFES"],
    [
        ["Fitness function terms", "4", "8"],
        ["Best fitness score", f"{mA['best_score']:.4f}", f"{mB['best_score']:.4f}"],
        ["Score inflation present", "Yes — length artefact", "No"],
        ["Best sequence length", f"{mA['best_seq_len']} aa ⚠", f"{mB['best_seq_len']} aa ✓"],
        ["Final pop. mean length", f"{mA['final_mean_len']:.0f} aa ⚠", f"{mB['final_mean_len']:.0f} aa ✓"],
        ["Best AMP probability", f"{mA['best_amp']:.4f}", f"{mB['best_amp']:.4f}"],
        ["Best hemolytic proxy", f"{mA['best_hemo']:.4f}", f"{mB['best_hemo']:.4f}"],
        ["Best pLDDT", f"{mA['best_plddt']:.4f}", f"{mB['best_plddt']:.4f}"],
        ["Best pTM", f"{mA['best_ptm']:.4f}", f"{mB['best_ptm']:.4f}"],
        ["Final mean score", f"{mA['final_mean_score']:.4f}", f"{mB['final_mean_score']:.4f}"],
        ["Drug candidate potential", "No — too long", "Yes — ~26 aa ✓"],
        ["AMP saturation generation", "~80", "~80"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. CONCLUSIONS
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Conclusions")

body("1.  Both pipelines work.", bold=True)
body("From random 24 aa starts, both converge to AMP-like sequences within 300 generations. The evolutionary framework is effective — pLDDT, pTM, and AMP probability all converge to high values regardless of fitness function design.")
doc.add_paragraph()

body("2.  Structural penalties are essential for drug-relevant candidates.", bold=True)
body("Without length and structural constraints, the optimizer exploits the duplication operator to produce 80 aa sequences that score well on the fitness function but are pharmacologically unrealistic. The length artefact is not a failure of the algorithm — it is a correct response to an incentive structure that rewards long sequences.")
doc.add_paragraph()

body("3.  Lower score ≠ worse biology.", bold=True)
body("A score of 0.595 satisfying 8 constraints (including length and helix) is more meaningful than 0.626 driven by length inflation. The MACREL+ESM3+PFES sequences are shorter, more consistent with natural AMPs, and within the therapeutic synthesis window.")
doc.add_paragraph()

body("4.  AMP classification saturates early.", bold=True)
body("MACREL stops discriminating after ~80 generations in both runs. The binary classifier reaches saturation and provides no further fitness gradient. This is a fundamental limitation: the signal strength collapses precisely when it would be most useful — during the plateau phase when structural optimization matters most.")
doc.add_paragraph()

body("5.  Hemolytic prediction remains open.", bold=True)
body("The biophysical proxy is a useful heuristic but cannot replace experimental validation. It correctly identifies strongly hemolytic sequences (mastoparan, melittin) but may misclassify borderline cases. All final candidates require in vitro hemolysis assay before any in vivo use.")
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 9. RECOMMENDED NEXT STEPS
# ══════════════════════════════════════════════════════════════════════════════
h1("9. Recommended Next Steps")

body("Computational:")
bullet("Integrate LLAMP: replace MACREL's binary AMP score with a quantitative log₁₀(MIC) prediction (μM). LLAMP provides a continuous signal throughout all 300 generations and would prevent AMP saturation.")
bullet("Run extended evolution (500–1,000 generations) with MACREL+ESM3+PFES to explore whether better sequences exist in the constrained space.")
bullet("Apply PFES to known AMP scaffolds (magainin-2, LL-37) as starting sequences rather than random sequences.")

body("Experimental (for top 5–10 MACREL+ESM3+PFES candidates):")
bullet("Synthesise peptides: ~€150/mg per 26 aa peptide (Fmoc solid-phase synthesis)")
bullet("MIC assay: E. coli K12, S. aureus ATCC 29213, P. aeruginosa PAO1")
bullet("Hemolysis assay: 0.5% human RBC, serial dilution 1–200 μg/mL")
bullet("CD spectroscopy: 50% TFE or SDS micelles — confirm amphipathic alpha-helix in membrane-mimetic environment")

# ══════════════════════════════════════════════════════════════════════════════
# 10. TECHNICAL DETAILS
# ══════════════════════════════════════════════════════════════════════════════
h1("10. Technical Details")

add_table(
    ["Component", "Details"],
    [
        ["Python", "3.11"],
        ["PyTorch", "2.6 with Apple MPS acceleration"],
        ["Structure predictor", "ESM3 (EvolutionaryScale, 2024)"],
        ["AMP classifier", "MACREL v1.2 (bioconda)"],
        ["Secondary structure", "PSIQUE"],
        ["Analysis", "visual_pfes.py"],
        ["Hardware", "Apple MacBook Air M2 (16 GB)"],
        ["Runtime per run", "~8–10 hours for 300 generations"],
        ["Repository branch (Run A)", "fitness-macrel"],
        ["Repository branch (Run B)", "fitness-macrel-pfes"],
        ["Original PFES branch", "alpha (unmodified)"],
    ]
)

body("Repository structure:")
bullet("fitness-macrel: MACREL+ESM3 run. Contains results_macrel_300/ with progress.log and analysis/.")
bullet("fitness-macrel-pfes: MACREL+ESM3+PFES run. Contains results_macrel_pfes_300/ with progress.log and analysis/.")
bullet("alpha: original PFES framework (no AMP scoring). Kept clean — no modifications.")

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
doc.save(OUT)
kb = os.path.getsize(OUT) // 1024
print(f"Saved → {OUT}  ({kb} KB)")
